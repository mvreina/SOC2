from html.parser import HTMLParser
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import logout
from .user import CustomUserCreationForm
from django.contrib.auth import authenticate, login
from .models import Project, Answer, Question, ProjectQuestion, Policy, ProjectPolicy, Text
from django.http import  JsonResponse
#from django.http import HttpResponse
from django.views.generic import CreateView, UpdateView, ListView
from .forms import ProjectForm, TextForm, TextProjectPolicyForm
from django.urls import reverse_lazy, reverse
from django.contrib import messages
import math
#User registration
from django.contrib.sites.shortcuts import get_current_site
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes
from django.template.loader import render_to_string
from django.contrib.auth.tokens import default_token_generator
from django.contrib.auth.models import User
from django.core.mail import EmailMessage
#Activate
from django.contrib.auth.tokens import default_token_generator
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_str
#Document
from django.http import HttpResponse
from docx import Document
import io
import pypandoc
import tempfile
import os
import mammoth
from bs4 import BeautifulSoup
from htmldocx import HtmlToDocx
import aspose.words as aw
from docx import Document

# Create your views here.
def home(request):
    return render(request, 'core/home.html')

# PROYECTOS
@login_required
def projects(request):
    # projectsList = Project.objects.all()
    # projectsList = Project.objects.all().order_by('name')
    # return render(request, 'core/projects.html' , {'projects': projectsList})
    return render(request, 'core/projects.html')

# Usado para construir la DataTable en la vista de lista
@login_required
def projectsList(request):
    # Filter by Current User
    projectsUser = Project.objects.filter(createdBy=request.user)
    projects = list(projectsUser.values())
    data = {'projects': projects}
    return JsonResponse(data)

# Ver proyecto
@login_required
def projectRead(request, pk):
    project = Project.objects.get(id=pk)
    projectPolicies = ProjectPolicy.objects.filter(project=project).order_by('orderProjectPolicy')
    totalPolicies = 32
    excludedPolicies = [int(value) for value in project.excludedPolicies]
    print(excludedPolicies)
    print(len(excludedPolicies))
    progressPolicies = math.ceil(((totalPolicies - len(excludedPolicies)) / totalPolicies) * 100)
    return render(request, 'core/projectRead.html', {'project': project, 'progressPolicies': progressPolicies, 'projectPolicies': projectPolicies})

class ProjectUpdateView(UpdateView):
    model = Project
    form_class = ProjectForm
    template_name = 'core/projectUpdate.html'
    success_url = reverse_lazy('projects')

    def form_valid(self, form):
        # Set the author to the currently logged-in user
        form.instance.createdBy = self.request.user
        messages.success(self.request, 'Proyecto actualizado correctamente.')
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, 'Error al actualizar proyecto.')
        return self.render_to_response(self.get_context_data(form=form))
    


# Crear proyecto
class ProjectCreateView(CreateView):

    model = Project
    form_class = ProjectForm
    template_name = 'core/projectCreate.html'
    success_url = reverse_lazy('projectQuestions')

    def form_valid(self, form):
        # Set the author to the currently logged-in user
        form.instance.createdBy = self.request.user
        project = form.save()
        #self.success_url = '/projectQuestions/' + str(project.id)
        self.success_url = reverse('projectQuestions', kwargs={'pk': project.id})
        # Create questions
        self.createQuestions(project)
        #messages.success(self.request, 'Proyecto creado correctamente.')   
        return super().form_valid(form) 

    def form_invalid(self, form):
        messages.warning(self.request, 'No se guardó el proyecto.')
        print(form.errors)
        return self.render_to_response(self.get_context_data(form=form))

    # Crear ProjectQuestion - Es la relación entre Proyecto y Preguntas y Respuestas
    def createQuestions(self, project):
        questions = Question.objects.all()
        for question in questions:
            answers = Answer.objects.filter(question=question)
            for answer in answers:
                if answer.type == 'checkbox':
                    userAnswer = True
                else:
                    userAnswer = False
                # Create ProjectQuestion
                projectQuestion = ProjectQuestion.objects.create(project=project, userAnswer=userAnswer, orderProjectQuestion=answer.orderAnswer)
                projectQuestion.question.add(question)
                projectQuestion.answer.add(answer)

# PREGUNTAS
class QuestionCreateView(CreateView):
    model = Question


# RESPUESTAS
class AnswerCreateView(CreateView):
    model = Answer

# PROYECTO - PREGUNTAS - RESPUESTAS
@login_required
def projectQuestions(request, pk):
    try:
        project = Project.objects.get(id=pk)
        question = Question.objects.get(orderQuestion=project.numQuestion)
        answers = Answer.objects.filter(question=question)
        totalQuestions = 13
        percentageProgress = math.ceil((project.numQuestion / totalQuestions) * 100)
        print("Pregunta actual:", project.numQuestion)

        if request.method == 'POST':
            #saveAnswers = request.POST.get("projectQuestions")
            #print(saveAnswers)
            projectQuestions = ProjectQuestion.objects.filter(project=project, question=question).order_by('orderProjectQuestion')
            projectQuestionId = []
            numSelected = 0 #Contador de respuestas seleccionadas    
            #Guardar respuestas tipo checkbox
            if question.type == 'checkbox':
                #print("Actualizando proyecto con preguntas tipo checkbox")
                for x in projectQuestions:
                    projectQuestionId.append(request.POST.get(str(x.id)))
                    #print(str(x.id))
                    #print(request.POST.get(str(x.id)))
                    if request.POST.get(str(x.id)) == None: 
                        x.userAnswer = False 
                    else: 
                        x.userAnswer = True
                        numSelected += 1
                        # Se excluyen la politicas dependiendo de las respuestas seleccionadas
                        #print("Respuesta tipo checkbox seleccionada política excluida:", x.answer.values()[0]['excludedPolicies'])
                        # Si la politica ya había sido excluida, no se vuelve agregar al proyecto
                        excludedPolicies = [int(value) for value in project.excludedPolicies]
                        answerExcludedPolicies = [int(value) for value in x.answer.values()[0]['excludedPolicies']]
                        print(excludedPolicies)
                        print(answerExcludedPolicies)
                        for expl in answerExcludedPolicies:
                            if not expl in excludedPolicies:
                                project.excludedPolicies += (expl,)
                                project.save()
                    x.save()
                    #print("Se guardó correctamente el proyectQuestion")
                #Si no ha seleccionado nada se muestra un mensaje de información y sigue en la misma pregunta
                if numSelected == 0:
                    messages.warning(request, "Por favor, seleccione una respuesta.")
                    context = {'projectQuestions': projectQuestions, 'question': question, 'percentageProgress': percentageProgress, 
                    'totalQuestions': totalQuestions, 'answers': answers, 'project': project}
                    return render(request, 'core/projectQuestions.html', context)
            else:
                #Guardar respuestas tipo radio
                #print("Actualizando proyecto con preguntas tipo radio")
                selected = request.POST.get('flexRadioDefault')
                #print("selected", selected)
                #print("Type", type(selected))
                #Si no ha seleccionado nada se muestra un mensaje de información y sigue en la misma pregunta
                if selected == None:
                    messages.warning(request, "Por favor, seleccione una respuesta.")
                    context = {'projectQuestions': projectQuestions, 'question': question, 'percentageProgress': percentageProgress, 
                    'totalQuestions': totalQuestions, 'answers': answers, 'project': project}
                    return render(request, 'core/projectQuestions.html', context)
                else:
                    #Si seleccionó una respuesta se guarda
                    for x in projectQuestions:
                        projectQuestionId.append(request.POST.get(str(x.id)))
                        #print(x.id)
                        #print(type(x.id))
                        if str(x.id) == selected: 
                            x.userAnswer = True
                            # Se excluyen la politicas dependiendo de las respuestas seleccionadas
                            #print("Respuesta tipo radio seleccionada política excluida:", x.answer.values()[0]['excludedPolicies'])
                            # Si la politica ya había sido excluida, no se vuelve agregar al proyecto
                            excludedPolicies = [int(value) for value in project.excludedPolicies]
                            answerExcludedPolicies = [int(value) for value in x.answer.values()[0]['excludedPolicies']]
                            print(excludedPolicies)
                            print(answerExcludedPolicies)
                            for expl in answerExcludedPolicies:
                                if not expl in excludedPolicies:
                                    project.excludedPolicies += (expl,)
                                    project.save()
                        else: 
                            x.userAnswer = False
                        print("Se guardó correctamente el proyecto")
                        x.save()

            print("projectQuestionId", projectQuestionId)


            #Avanzar a la siguiente respuesta
            if 'finish' in request.POST:
                project.numQuestion += 1
                project.save()
                # Crea las politicas y las asocia al proyecto
                createPolicies(project)
                success_url = reverse('projectRead', kwargs={'pk': project.id})
                return redirect(success_url)
            elif 'next' in request.POST:
                
                project.numQuestion += 1
                project.save()
                question = Question.objects.get(orderQuestion=project.numQuestion)
                answers = Answer.objects.filter(question=question)
                projectQuestions = ProjectQuestion.objects.filter(project=project, question=question).order_by('orderProjectQuestion')
                percentageProgress = math.ceil((project.numQuestion / totalQuestions) * 100)
                print("Avanza a la pregunta: " + str(project.numQuestion))


        else:
            #projectQuestions = ProjectQuestion.objects.filter(project=project, question=question)
            #projectQuestions = ProjectQuestionForm(initial={'project': project, 'question': question})
            projectQuestions = ProjectQuestion.objects.filter(project=project, question=question).order_by('orderProjectQuestion')

        #print(projectQuestions)
        #print(projectQuestions.first())
        #print(projectQuestions.first().answer.values()[0]['text'])
        #print(list(projectQuestions))
        #list_projectQuestion = list(projectQuestions)
        #for x in list_projectQuestion:
            #print(x.answer.values()[0]['text'])

        #for x in projectQuestions:
        #    print(x.answer.values()[0]['text'])
        #print(answers.values())
        context = {'projectQuestions': projectQuestions, 'question': question, 'percentageProgress': percentageProgress, 
                    'totalQuestions': totalQuestions, 'answers': answers, 'project': project}
        return render(request, 'core/projectQuestions.html', context)
      
    except Exception as e:
        messages.warning(request, "Se produjo una excepción, consulte al Administrador")
        print("Se produjo una excepción: "+str(e))
        return redirect('projects')

# Crea las politicas asociadas al proyecto
def createPolicies(project):
    policies = Policy.objects.all()
    policyDocuments = Text.objects.all()
    for policy in policies:
        
        #print("Creando las políticas del proyecto:", project.excludedPolicies, " - ", policy.orderPolicy)
        excludedPolicies = [int(value) for value in project.excludedPolicies]
        
        policyText = policyDocuments.get(orderText=policy.orderPolicy)

        # Si la política no está en la lista de excluidas el flag excluded es False, sino es True
        if policy.orderPolicy in excludedPolicies:
            ProjectPolicy.objects.create(project=project, policy=policy, orderProjectPolicy=policy.orderPolicy, excluded=True, content=policyText.content, fileName=policyText.fileName)
        else:
            ProjectPolicy.objects.create(project=project, policy=policy, orderProjectPolicy=policy.orderPolicy, excluded=False, content=policyText.content, fileName=policyText.fileName)

class ProjectListView(ListView):
    model = Project

class AnswerListView(ListView):
    model = Answer

class QuestionListView(ListView):
    model = Question



### DOCUMENTOS
"""
def edit_text(request):
    if request.method == 'POST':
        form = TextForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('edit_text')
    else:
        form = TextForm()

    return render(request, 'core/textUpdate.html', {'form': form})

    def edit_text(request, text_id):
    text = get_object_or_404(Text, pk=text_id)
    
    if request.method == 'POST':
            form = TextForm(request.POST)
            if form.is_valid():
                form.save()
                form = TextForm(initial={'content': text.content, 'name': text.name})
                return render(request, 'core/textUpdate.html', {'form': form})
            else:
                form = TextForm(initial={'form': form})

    else:
        form = TextForm(initial={'content': text.content, 'name': text.name})

    return render(request, 'core/textUpdate.html', {'form': form})
"""




class TextUpdateView(UpdateView):
    model = Text
    form_class = TextForm
    template_name = 'core/textUpdate.html'

    def form_valid(self, form):
        #modify_docx_document(form.instance.id)
        messages.success(self.request, 'Documento actualizado correctamente.')
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, 'Error al actualizar documento.')
        print(form.errors)
        return self.render_to_response(self.get_context_data(form=form))
    
    def get_success_url(self):
        return reverse('textUpdate', args=[self.object.pk])


class TextProjectPolicyUpdateView(UpdateView):
    model = ProjectPolicy
    form_class = TextProjectPolicyForm
    template_name = 'core/textProjectPolicyUpdate.html'

    def form_valid(self, form):
        #modify_docx_document(form.instance.id)
        messages.success(self.request, 'Documento actualizado correctamente.')
        return super().form_valid(form)

    def form_invalid(self, form):
        messages.error(self.request, 'Error al actualizar documento.')
        print(form.errors)
        return self.render_to_response(self.get_context_data(form=form))
    
    def get_success_url(self):
        return reverse('textProjectPolicyUpdate', args=[self.object.pk])

def download_word_document(request, text_id):
    
    text = get_object_or_404(Text, pk=text_id)


    
    #htmlToDocx3()




    #docx_document = html_to_docx(text.content)
    html_content = "<p>This is a simple HTML to DOCX conversion.</p>"
    output_path = "output.docx"
    html_to_docx(html_content, output_path)


    # Specify the output file path
    #output_file_path = "WorkingFromHome.docx"

    # Save the modified DOCX document to the output file
    #docx_document.save(output_file_path)

    # Create an in-memory buffer to store the Word document
    #buffer = io.BytesIO()
    #docx_document.save(buffer)
    #buffer.seek(0)


    #response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    #response = HttpResponse(buffer.read(), content_type='application/msword')
    #response['Content-Disposition'] = f'attachment; filename={output_file_path}'


    #return response


def html_to_docx(html_text):
    # Create a temporary output file
    with tempfile.NamedTemporaryFile(suffix=".md", delete=False) as tmp_file:
        tmp_file_path1 = tmp_file.name

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
        tmp_file_path2 = tmp_file.name

    # Convert HTML to DOCX and save it to the temporary file
    #pypandoc.convert_text(html_text, 'md', format='html', outputfile=tmp_file_path1)

    #output = pypandoc.convert_file(tmp_file_path1, 'docx', format='md', outputfile=tmp_file_path2)

    pypandoc.convert_text(html_text, 'docx', format='html', outputfile=tmp_file_path2)

    #pypandoc.convert_text(html_text, 'docx', format='html', outputfile=tmp_file_path)

    # Open the temporary DOCX file
    #docx_document = Document(tmp_file_path)

    new_parser = HtmlToDocx()
    new_parser.parse_html_file("text.html", "text.docx")

    docx_document = Document(tmp_file_path2)
   

    # Clean up the temporary file
    #os.remove(tmp_file_path)

    return docx_document


def htmlToDocx(html_text):
    # Create a new DOCX document
    new_doc = Document()

    # Use BeautifulSoup to extract plain text from the HTML
    soup = BeautifulSoup(html_text, 'html.parser')
    plain_text = soup.get_text()

    # Split the plain text content by the page break marker
    text_content_parts = plain_text.split('SALTODELINEA')

    # Iterate through the plain text content parts and add them to the DOCX document
    for part in text_content_parts:
        if part.strip():
            new_doc.add_paragraph(part)
            new_doc.add_page_break()  # Add a page break after each part
    
    return new_doc

def modify_docx_document(text_id):
    text = get_object_or_404(Text, pk=text_id)

    # Convert HTML to DOCX document
    docx_document = html_to_docx(text.content)

    # Specify the output file path
    output_file_path = "modified_word_document.docx"

    # Save the modified DOCX document to the output file
    docx_document.save(output_file_path)

    # Provide the output file for download
    with open(output_file_path, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={output_file_path}'

def htmToDocx2():

    html_content = "<h1>This is an example HTML content.</h1>"
    preHtml = "<html xmlns:o='urn:schemas-microsoft-com:office:office' xmlns:w='urn:schemas-microsoft-com:office:word' xmlns='http://www.w3.org/TR/REC-html40'>"
    postHtml = "</body></html>"
    
    html = preHtml + html_content + postHtml

    # Convert HTML to Word
    result = mammoth.convert(html_content)

    # Prepare the response with the Word document
    response = HttpResponse(content_type='application/msword')
    response['Content-Disposition'] = 'attachment; filename=converted.docx'
    response.write(result.value)


def htmlToDocx3():
    # Load the HTML file from disk
    doc = aw.Document("text.html")

    # Save the HTML file as Word DOCX document
    doc.save("html-to-word.docx")




class HTMLToDocxConverter(HTMLParser):
    def __init__(self):
        super().__init__()
        self.doc = Document()
        self.current_paragraph = None

    def handle_starttag(self, tag, attrs):
        if tag == 'p':
            self.current_paragraph = self.doc.add_paragraph()

    def handle_endtag(self, tag):
        if tag == 'p':
            self.current_paragraph = None

    def handle_data(self, data):
        if self.current_paragraph:
            self.current_paragraph.add_run(data)

def html_to_docx(html_content, output_path):
    doc = Document()
    result = mammoth.convert(html_content)

    doc.add_paragraph(result.value)
    
    doc.save(output_path)

### AUTENTICACIÓN, REGISTRO, VALIDACIÓN DE USUARIOS

# Logout
def exit(request):
    logout(request)
    return redirect('home')


# User registration - queda deprecated - se usa ahora registerEmail
def register(request):

    data = {
        'form': CustomUserCreationForm()
    }

    if request.method == 'POST':
        user_creation_form = CustomUserCreationForm(data=request.POST)

        if user_creation_form.is_valid():
            user_creation_form.save()

            user = authenticate(
                username=user_creation_form.cleaned_data['username'],
                password=user_creation_form.cleaned_data['password1']
            )

            login(request, user)

            return redirect('home')
        else:
            data['form'] = user_creation_form
            return render(request, 'registration/register.html', data)


    return render(request, 'registration/register.html', data)


def registerEmail(request):

    data = {
        'form': CustomUserCreationForm()
    }

    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        print("Entra a register_email")
        if form.is_valid():
            print("Entra a register_email - is_valid")
            user = form.save()
            user.is_active = False  # Mark the user as inactive until they confirm their email
            user.save()
            current_site = get_current_site(request)
            mail_subject = 'Active su cuenta en M2J'
            if request.is_secure():
                protocol = "https"
            else:
                protocol = "http"

            message = render_to_string(
                'registration/activationEmail.html',
                {
                    'user': user,
                    'domain': current_site.domain,
                    'uid': urlsafe_base64_encode(force_bytes(user.pk)),
                    'token': default_token_generator.make_token(user),
                    'protocol': protocol
                },
            )
            print(message)
            email = EmailMessage(mail_subject, message, to=[user.email])
            email.content_subtype = 'html'
            emailsend = email.send()
            print(emailsend)
            messages.info(request, 'Envió correo electrónico para activar su cuenta.')
            #Allow access to registration success page
            request.session['allowed_access'] = True
            return redirect('registrationSuccess')  # Create this view
        else:
            print("Entra a register_email - is_invalid")
            data['form'] = form
            return render(request, 'registration/register.html', data)
            
    return render(request, 'registration/register.html', data)


def activate(request, uidb64, token):
    try:
        print("Entra a activate")
        uid = force_str(urlsafe_base64_decode(uidb64))
        print("uid", uid)
        user = User.objects.get(pk=uid)
        print("user", user)
        if default_token_generator.check_token(user, token):
            user.is_active = True
            user.save()
            #login(request, user)
            request.session['allowed_access'] = True
            return redirect('activationSuccess')
        else:
            return redirect('activationFailure')
    except Exception as e:
        print("Se produjo una excepción: "+str(e))
        request.session['allowed_access'] = False
    return render(request, 'registration/activationInvalid.html')

def activationSuccess(request):
    return render(request, 'registration/activationSuccess.html')

def activationFailure(request):
    logout(request)
    return render(request, 'registration/activationFailure.html')

def registrationSuccess(request):
    return render(request, 'registration/registrationSuccess.html')

def activationInvalid(request):
    logout(request)
    return render(request, 'registration/activationInvalid.html')

def accessDenied(request):
    return render(request, 'registration/accessDenied.html')

#404 Not Found
def pageNotFound(request):
    return render(request, 'core/404.html', status=404)

# OTHERS
@login_required
def aboutUs(request):
    # projectsList = Project.objects.all()
    # projectsList = Project.objects.all().order_by('name')
    # return render(request, 'core/projects.html' , {'projects': projectsList})
    return render(request, 'core/aboutUs.html')